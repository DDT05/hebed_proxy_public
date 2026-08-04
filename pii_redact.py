"""
pii_redact.py --- mitmproxy addon for PII anonymization.

Architecture (synced with anonymize.py):
  - NlpEngineProvider + spaCy for multi-language NER (en/fr/de/it)
  - Graceful fallback to NoOpNlpEngine if spaCy models not installed
  - All 12 custom recognizers from the reference implementation
  - Custom placeholder-based redaction: reversible for response restoration
  - AnonymizerEngine imported and available for direct-use mode

Run: mitmdump --listen-port 8080 -s pii_redact.py
"""

import io
import json
import os
import re
import sys
import threading
from datetime import datetime, timezone
from pathlib import Path

# Fix: Python313 user site-packages may not be on sys.path when the base
# Python install is missing (mitmdump uses a pip launcher that resolves to a
# Python that may have been relocated).  Inject it so Presidio + spaCy import.
_site = os.path.join(os.environ.get("APPDATA", ""), "Python", "Python313", "site-packages")
if os.path.isdir(_site) and _site not in sys.path:
    sys.path.insert(0, _site)

from mitmproxy import http
from urllib.parse import parse_qs, urlencode, unquote
from email.parser import BytesParser

# ---------------------------------------------------------------------------
# Language: default "en", can be overridden. Custom regex recognizers are
# language-agnostic; the language setting affects spaCy NER only.
# ---------------------------------------------------------------------------
DEFAULT_LANGUAGE = "en"

# ---------------------------------------------------------------------------
# Log helpers --- structured PII events (consumed by proxy_mvp Tauri app)
# ---------------------------------------------------------------------------
_log_lock = threading.Lock()
_log_file: str | None = None


def _log_path():
    global _log_file
    if _log_file is None:
        log_dir = os.environ.get(
            "PII_LOG_DIR",
            str(Path(__file__).resolve().parent),
        )
        _log_file = os.path.join(log_dir, "pii_events.log")
    return _log_file


def _log(event: dict):
    """Append a structured PII event to pii_events.log."""
    entry = {"ts": datetime.now(timezone.utc).isoformat(), **event}
    try:
        with _log_lock:
            with open(_log_path(), "a", encoding="utf-8") as f:
                f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except Exception:
        pass  # never crash the proxy for logging


# ---------------------------------------------------------------------------
# File upload log --- records extracted text + PII findings from file uploads
# ---------------------------------------------------------------------------
_files_log_file: str | None = None


def _files_log_path():
    global _files_log_file
    if _files_log_file is None:
        log_dir = os.environ.get(
            "PII_LOG_DIR",
            str(Path(__file__).resolve().parent),
        )
        _files_log_file = os.path.join(log_dir, "files.log")
    return _files_log_file


def _log_file_event(event: dict):
    """Append a file-scan event to files.log."""
    entry = {"ts": datetime.now(timezone.utc).isoformat(), **event}
    try:
        with _log_lock:
            with open(_files_log_path(), "a", encoding="utf-8") as f:
                f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Prompt interception log --- records original + transformed prompt content
# ---------------------------------------------------------------------------
_prompts_log_file: str | None = None


def _prompts_log_path():
    global _prompts_log_file
    if _prompts_log_file is None:
        log_dir = os.environ.get(
            "PII_LOG_DIR",
            str(Path(__file__).resolve().parent),
        )
        _prompts_log_file = os.path.join(log_dir, "prompts.log")
    return _prompts_log_file


def _log_prompt(event: dict):
    """Append an intercepted-prompt event (original + transformed) to prompts.log."""
    entry = {"ts": datetime.now(timezone.utc).isoformat(), **event}
    try:
        with _log_lock:
            with open(_prompts_log_path(), "a", encoding="utf-8") as f:
                f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except Exception:
        pass  # never crash the proxy for logging


# ---------------------------------------------------------------------------
# Presidio engine --- warmed in background at startup
# ---------------------------------------------------------------------------
_engine = None          # AnalyzerEngine ready when != None
_anonymizer = None      # AnonymizerEngine (available for direct-use mode)
_engine_ready = threading.Event()


def _init_engine():
    """Initialize Presidio AnalyzerEngine with multi-language NER.

    Tries NlpEngineProvider with spaCy first (en/fr/de/it NER).
    Falls back to NoOpNlpEngine (regex-only, instant) if models missing.
    """
    global _engine, _anonymizer

    try:
        from presidio_analyzer import (
            AnalyzerEngine,
            Pattern,
            PatternRecognizer,
        )
        from presidio_anonymizer import AnonymizerEngine

        # --- Try spaCy NlpEngineProvider first ---
        nlp_engine = None
        try:
            from presidio_analyzer.nlp_engine import NlpEngineProvider

            # Resolve config: Tauri copies addon into target/debug/ but not the YAML.
            # Walk up from addon: debug/ -> target/ -> src-tauri/ -> project root.
            addon_dir = Path(__file__).resolve().parent
            candidates = [
                addon_dir / "spacy_en_fr_de_it.yaml",
                addon_dir.parent / "spacy_en_fr_de_it.yaml",
                addon_dir.parent.parent / "spacy_en_fr_de_it.yaml",
                addon_dir.parent.parent.parent / "spacy_en_fr_de_it.yaml",
            ]
            config_path = next((c for c in candidates if c.exists()), None)
            if config_path:
                provider = NlpEngineProvider(conf_file=str(config_path))
                nlp_engine = provider.create_engine()
                print("[PII] spaCy NlpEngine loaded (en/fr/de/it NER)", file=sys.stderr, flush=True)
            else:
                print("[PII] spaCy config not found, falling back to NoOpNlpEngine",
                      file=sys.stderr, flush=True)
        except Exception as e:
            print(f"[PII] spaCy NlpEngine unavailable ({e}), falling back to NoOpNlpEngine",
                  file=sys.stderr, flush=True)

        # --- Fallback: NoOpNlpEngine (regex-only) ---
        if nlp_engine is None:
            from presidio_analyzer.nlp_engine import NoOpNlpEngine
            nlp_engine = NoOpNlpEngine(
                models=[{"lang_code": "en", "model_name": "no_op"}]
            )
            print("[PII] NoOpNlpEngine loaded (regex-only mode)", file=sys.stderr, flush=True)

        # --- Create engines ---
        _engine = AnalyzerEngine(
            nlp_engine=nlp_engine,
            supported_languages=["en", "fr", "de", "it"],
        )
        _anonymizer = AnonymizerEngine()

        # --- Custom recognizers (all 12 from reference implementation) ---
        _register_custom_recognizers(_engine)

        # --- Self-test: verify engine detects known PII ---
        _self_test(_engine)

        _engine_ready.set()
        print("[PII] engine ready", file=sys.stderr, flush=True)

    except Exception as e:
        print(f"[PII] engine init failed: {e}", file=sys.stderr, flush=True)
        _engine_ready.set()  # don't block forever


def _self_test(engine):
    """Run a quick self-test to verify PII detection works."""
    test_text = "my email is john@doe.com and phone number +4176573231311"
    try:
        results = engine.analyze(text=test_text, language="en")
        if results:
            found = [f"{r.entity_type}({r.score:.2f})" for r in results]
            print(f"[PII] self-test: {len(results)} match(es) -> {found}", file=sys.stderr, flush=True)
        else:
            print("[PII] self-test: 0 matches (ENGINE RETURNED EMPTY — recognizers may not be active)",
                  file=sys.stderr, flush=True)
    except Exception as e:
        print(f"[PII] self-test ERROR: {e}", file=sys.stderr, flush=True)
        import traceback
        traceback.print_exc(file=sys.stderr)


def _register_custom_recognizers(engine):
    """Register all custom regex recognizers --- synced with edit1_py.txt."""
    from presidio_analyzer import Pattern, PatternRecognizer

    recognizers = [
        # 1. IBAN
        PatternRecognizer(
            supported_entity="IBAN",
            name="iban",
            patterns=[Pattern(
                name="iban",
                regex=r"\b[A-Z]{2}\s?\d{2}\s?[A-Z0-9]{4}\s?[A-Z0-9]{4}\s?[A-Z0-9]{4}\s?[A-Z0-9]{1,8}\b",
                score=0.95,
            )],
        ),
        # 2. International phone numbers (+XX XXX XXX XXXX)
        PatternRecognizer(
            supported_entity="PHONE_NUMBER",
            name="phone_international",
            patterns=[Pattern(
                name="phone_intl",
                regex=r"\+\d{1,3}[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}\b",
                score=0.85,
            )],
        ),
        # 3. US phone numbers (XXX) XXX-XXXX
        PatternRecognizer(
            supported_entity="PHONE_NUMBER",
            name="phone_us",
            patterns=[Pattern(
                name="phone_us",
                regex=r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b",
                score=0.80,
            )],
        ),
        # 4. Email addresses
        PatternRecognizer(
            supported_entity="EMAIL_ADDRESS",
            name="email",
            patterns=[Pattern(
                name="email",
                regex=r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
                score=0.95,
            )],
        ),
        # 5. French phone numbers (0X XX XX XX XX or +33 X XX XX XX)
        PatternRecognizer(
            supported_entity="PHONE_NUMBER",
            name="phone_french",
            patterns=[Pattern(
                name="phone_fr",
                regex=r"\+?33[\s.-]?\d[\s.-]?\d{2}[\s.-]?\d{2}[\s.-]?\d{2}\b",
                score=0.80,
            )],
        ),
        # 6. Italian phone numbers (+39 XX XXX XXXX or 0XX XXX XXXX)
        PatternRecognizer(
            supported_entity="PHONE_NUMBER",
            name="phone_italian",
            patterns=[Pattern(
                name="phone_it",
                regex=r"\+?39[\s.-]?\d{2,3}[\s.-]?\d{3,4}[\s.-]?\d{3,4}\b",
                score=0.80,
            )],
        ),
        # 7. Italian fiscal code (Codice Fiscale)
        # Format: 6 letters + 2 digits + 1 letter + 4 digits + 1 letter + 2 digits + 1 letter
        PatternRecognizer(
            supported_entity="IT_FISCAL_CODE",
            name="it_fiscal_code",
            patterns=[Pattern(
                name="it_cf",
                regex=r"\b[A-Z]{6}\d{2}[A-Z]\d{3}[A-Z]\d{2}[A-Z]\b",
                score=0.70,
            )],
        ),
        # 8. Credit card numbers
        PatternRecognizer(
            supported_entity="CREDIT_CARD",
            name="credit_card",
            patterns=[Pattern(
                name="credit_card",
                regex=r"\b(?:\d{4}[-\s]?){3}\d{4}\b|\b(?:\d{4}[-\s]?){2}\d{4}[-\s]?\d{4}\b|\b\d{13,19}\b",
                score=0.85,
            )],
        ),
        # 9. French social security number (N° Sécurité Sociale)
        # Score intentionally low to avoid false positives in IBANs, phone numbers, etc.
        PatternRecognizer(
            supported_entity="FR_SSN",
            name="fr_ssn",
            patterns=[Pattern(
                name="fr_ssn",
                regex=r"(?<![A-Z0-9])(?<!\d)[12]\d{2}[0123456789]\d{2}[0-9]{2}\d{3}[0-9]{5}[0-9]{2}[0-9]{3}(?!\d)",
                score=0.50,
            )],
        ),
        # 10. German ID card (Personalausweis / Reisepass)
        PatternRecognizer(
            supported_entity="DE_ID",
            name="de_id",
            patterns=[Pattern(
                name="de_passport",
                regex=r"\bDE\d{7,9}\b",
                score=0.60,
            )],
        ),
        # 11. US SSN (XXX-XX-XXXX) --- retained from previous pii_redact.py
        PatternRecognizer(
            supported_entity="US_SSN",
            name="us_ssn",
            patterns=[Pattern(
                name="us_ssn",
                regex=r"\b\d{3}-\d{2}-\d{4}\b",
                score=0.85,
            )],
        )
    ]

    for rec in recognizers:
        engine.registry.add_recognizer(rec)


# Start background init
threading.Thread(target=_init_engine, daemon=True).start()

# ---------------------------------------------------------------------------
# PII store --- maps flow_id -> {placeholder: original_value}
# ---------------------------------------------------------------------------
_pii_store: dict[str, dict[str, str]] = {}


def _redact(text: str, language: str = DEFAULT_LANGUAGE) -> tuple[str, dict[str, str]]:
    """Detect PII and replace with numbered placeholders.

    Returns (redacted_text, {placeholder: original_value}).
    """
    if not text or not text.strip():
        return text, {}

    if _engine is None:
        # Wait up to 3 seconds for engine init to complete
        if _engine_ready.wait(timeout=3.0):
            pass  # engine is ready now, proceed
        else:
            print("[PII] engine not ready after 3s", file=sys.stderr, flush=True)
            return text, {}

    try:
        results = _engine.analyze(text=text, language=language)
    except Exception as e:
        print(f"[PII] analyze error: {e}", file=sys.stderr, flush=True)
        return text, {}

    if not results:
        return text, {}

    # Sort by start position, deduplicate overlapping spans
    results = sorted(results, key=lambda r: r.start)

    store: dict[str, str] = {}
    counters: dict[str, int] = {}
    parts: list[str] = []
    last_end = 0

    for r in results:
        if r.start < last_end:
            continue  # skip overlapping

        entity = r.entity_type
        idx = counters.get(entity, 0)
        counters[entity] = idx + 1

        placeholder = f"[{entity}_{idx:03d}]"
        store[placeholder] = text[r.start : r.end]

        parts.append(text[last_end : r.start])
        parts.append(placeholder)
        last_end = r.end

    parts.append(text[last_end:])
    redacted = "".join(parts)

    # Log structured event
    entities_found = [
        {"type": r.entity_type, "score": round(r.score, 3)}
        for r in results
    ]
    _log({
        "event": "redact",
        "entities": entities_found,
        "original_len": len(text),
        "redacted_len": len(redacted),
    })

    return redacted, store


# ---------------------------------------------------------------------------
# JSON body scanner --- traverses known LLM API request structures
# ---------------------------------------------------------------------------
def _scan_json(body: dict, url: str = "", content_type: str = "") -> tuple[bytes, int]:
    """Scan JSON body for PII in known LLM API fields. Returns (new_body, count)."""
    count = 0
    texts_found = 0
    samples: list[str] = []

    def _scan(text: str, label: str) -> tuple[str, dict[str, str]]:
        nonlocal texts_found, count
        texts_found += 1
        redacted, store = _redact(text)
        preview = text[:120].replace("\n", "\\n")
        _log_prompt({
            "event": "prompt_redact",
            "flow_id": flow_id(),
            "url": url[:300],
            "content_type": content_type[:80],
            "label": label,
            "original_len": len(text),
            "redacted_len": len(redacted),
            "pii_count": len(store),
            "original": text[:2000],
            "redacted": redacted[:2000],
        })
        if store:
            count += 1
            print(f"[PII] {label}: {preview}", flush=True)
        else:
            samples.append(f"{label}: {preview}")
        return redacted, store

    # OpenAI-style: messages[].content.parts[]
    for i, msg in enumerate(body.get("messages", [])):
        content = msg.get("content", {})
        parts = content.get("parts", [])
        new_parts = []
        for j, part in enumerate(parts):
            if isinstance(part, str) and part.strip():
                redacted, store = _scan(part, f"msg[{i}].parts[{j}]")
                if store:
                    _pii_store[flow_id()] = store
                new_parts.append(redacted)
            else:
                new_parts.append(part)
        if new_parts != parts:
            content["parts"] = new_parts

        # Also check if content itself is a string
        c = msg.get("content")
        if isinstance(c, str) and c.strip():
            redacted, store = _scan(c, f"msg[{i}].content")
            if store:
                _pii_store[flow_id()] = store
                msg["content"] = redacted

    # Direct prompt field
    if isinstance(body.get("prompt"), str) and body["prompt"].strip():
        redacted, store = _scan(body["prompt"], "prompt")
        if store:
            _pii_store[flow_id()] = store
            body["prompt"] = redacted

    if texts_found > 0 and count == 0:
        print(f"[PII] scanned {texts_found} field(s), no PII:", file=sys.stderr, flush=True)
        for s in samples:
            print(f"  {s}", file=sys.stderr, flush=True)

    return json.dumps(body).encode(), count


# ---------------------------------------------------------------------------
# Form-data scanner --- handles application/x-www-form-urlencoded endpoints
# ---------------------------------------------------------------------------
def _scan_form(data: bytes, url: str = "", content_type: str = "") -> tuple[bytes, int]:
    """Scan URL-encoded form data for PII in the 'prompt' field.
    Returns (new_body_bytes, count)."""
    text = data.decode("utf-8", errors="replace")
    params = parse_qs(text, keep_blank_values=True)

    count = 0
    prompt_vals = params.get("prompt", [])
    new_vals = []
    for val in prompt_vals:
        decoded = unquote(val)
        if decoded.strip():
            redacted, store = _redact(decoded)
            _log_prompt({
                "event": "prompt_redact",
                "flow_id": flow_id(),
                "url": url[:300],
                "content_type": content_type[:80],
                "label": "form.prompt",
                "original_len": len(decoded),
                "redacted_len": len(redacted),
                "pii_count": len(store),
                "original": decoded[:2000],
                "redacted": redacted[:2000],
            })
            if store:
                _pii_store[flow_id()] = store
                count += 1
                new_vals.append(redacted)
                print(f"[PII] form.prompt: {decoded[:120]}", flush=True)
            else:
                new_vals.append(decoded)
                print(f"[PII] form.prompt (no PII): {decoded[:120]}", file=sys.stderr, flush=True)
        else:
            new_vals.append(decoded)

    if count > 0:
        params["prompt"] = new_vals
        # Rebuild: urlencode preserves the original parameter order well enough
        new_text = urlencode(params, doseq=True)
        return new_text.encode("utf-8"), count

    return data, 0


# ---------------------------------------------------------------------------
# File extraction helpers
# ---------------------------------------------------------------------------
def _extract_text_from_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def _extract_text_from_pdf(data: bytes) -> str:
    try:
        import fitz
    except ImportError:
        return "[PyMuPDF not installed]"
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        parts = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(parts)
    except Exception as e:
        return f"[PDF error: {e}]"


def _extract_text_from_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return "[python-docx not installed]"
    try:
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        return f"[DOCX error: {e}]"


def _extract_text_from_image(data: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return "[pytesseract/PIL not installed]"
    try:
        img = Image.open(io.BytesIO(data))
        try:
            text = pytesseract.image_to_string(img, lang="eng+fra+deu")
        except Exception:
            text = pytesseract.image_to_string(img, lang="eng")
        return text.strip()
    except Exception as e:
        return f"[OCR error: {e}]"


def _sniff_content_type(data: bytes) -> str:
    """Detect content type from magic bytes when the header is missing/empty."""
    if data[:4] == b"%PDF":
        return "application/pdf"
    if data[:3] == b"\xff\xd8\xff":
        return "image/jpeg"
    if data[:4] == b"\x89PNG":
        return "image/png"
    if data[:4] == b"GIF8":
        return "image/gif"
    if data[:2] == b"PK":
        # docx/xlsx/pptx are zip containers; default to docx-family marker
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return "image/webp"
    try:
        data.decode("utf-8")
        return "text/plain"
    except Exception:
        return "application/octet-stream"


def _extract_text(data: bytes, content_type: str) -> str:
    ct = content_type.lower()
    if "text/plain" in ct:
        return _extract_text_from_txt(data)
    if "application/pdf" in ct:
        return _extract_text_from_pdf(data)
    if "wordprocessingml" in ct or "msword" in ct:
        return _extract_text_from_docx(data)
    if any(t in ct for t in ("image/jpeg", "image/png", "image/jpg", "image/webp", "image/bmp", "image/tiff")):
        return _extract_text_from_image(data)
    # Fallback: sniff magic bytes when the header is empty or generic
    if not ct or ct in ("application/octet-stream", "binary/octet-stream"):
        sniffed = _sniff_content_type(data)
        if sniffed != ct:
            return _extract_text(data, sniffed)
    return f"[unsupported: {content_type}]"


# ---------------------------------------------------------------------------
# File redaction helpers (only active when PII_REDACT_FILES=1)
# ---------------------------------------------------------------------------
def _redact_pdf(data: bytes, store: dict) -> bytes:
    """Redact PII from PDF using PyMuPDF redaction API. Text is permanently removed."""
    try:
        import fitz
    except ImportError:
        return data
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        for page in doc:
            for original in store.values():
                areas = page.search_for(original)
                for area in areas:
                    page.add_redact_annot(area, fill=(0, 0, 0))
            page.apply_redactions()
        out = doc.tobytes(garbage=4, deflate=True)
        doc.close()
        return out
    except Exception as e:
        print(f"[PII] PDF redaction failed: {e}", file=sys.stderr, flush=True)
        return data


def _redact_image(data: bytes, store: dict) -> bytes:
    """Draw black rectangles over PII text regions in images."""
    try:
        import pytesseract
        from PIL import Image, ImageDraw
    except ImportError:
        return data
    try:
        img = Image.open(io.BytesIO(data))
        ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        draw = ImageDraw.Draw(img)
        for original in store.values():
            for i, word in enumerate(ocr_data["text"]):
                if word.strip() and word.strip() in original:
                    x, y = ocr_data["left"][i], ocr_data["top"][i]
                    w, h = ocr_data["width"][i], ocr_data["height"][i]
                    draw.rectangle([x, y, x + w, y + h], fill="black")
        out = io.BytesIO()
        img.save(out, format=img.format or "PNG")
        return out.getvalue()
    except Exception as e:
        print(f"[PII] Image redaction failed: {e}", file=sys.stderr, flush=True)
        return data


def _redact_docx(data: bytes, store: dict) -> bytes:
    """Replace PII text with placeholders in DOCX runs."""
    try:
        from docx import Document
    except ImportError:
        return data
    try:
        doc = Document(io.BytesIO(data))
        for para in doc.paragraphs:
            for run in para.runs:
                for placeholder, original in store.items():
                    run.text = run.text.replace(original, placeholder)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for placeholder, original in store.items():
                                run.text = run.text.replace(original, placeholder)
        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception as e:
        print(f"[PII] DOCX redaction failed: {e}", file=sys.stderr, flush=True)
        return data

# ---------------------------------------------------------------------------
# File upload scanner — intercept, extract, scan, optionally redact
# Supports batch (multiple files in multipart) and single-file uploads.
# ---------------------------------------------------------------------------
def _scan_file_upload(flow: http.HTTPFlow) -> tuple[bytes, int]:
    """Extract text from all uploaded files, scan each for PII, log individually.

    Returns (body_bytes, total_pii_count).
    """
    content_type = flow.request.headers.get("content-type", "")
    data = flow.request.content
    url = flow.request.pretty_url
    total_pii = 0
    file_index = 0

    # --- Multipart: collect all file parts ---
    if "multipart/form-data" in content_type:
        # Parse boundary
        boundary = None
        for part in content_type.split(";"):
            p = part.strip()
            if p.startswith("boundary="):
                boundary = p.split("=", 1)[1].strip("\"'")
                break

        if not boundary:
            print(f"[PII] file: {flow.id} no boundary in multipart", file=sys.stderr, flush=True)
            return data, 0

        try:
            header = "Content-Type: " + content_type + "\r\n\r\n"
            msg = BytesParser().parsebytes(header.encode() + data)
            parts = list(msg.walk())
            modified_parts: dict[int, bytes] = {}  # index -> new payload
            file_count = 0

            for idx, part_msg in enumerate(parts):
                disp = str(part_msg.get("Content-Disposition", ""))
                if "filename=" not in disp:
                    continue

                file_count += 1
                file_data = part_msg.get_payload(decode=True)
                file_ct = part_msg.get_content_type()
                file_size = len(file_data)
                filename = ""
                for token in disp.split(";"):
                    token = token.strip()
                    if token.startswith("filename="):
                        filename = token.split("=", 1)[1].strip("\"'")
                        break

                print(f"[PII] multipart[{file_count}]: {flow.id} filename={filename} ct={file_ct} {file_size}B",
                      file=sys.stderr, flush=True)

                # Process this file
                pii, new_data = _process_one_file(
                    flow.id, file_data, file_ct, url, filename, file_count
                )
                total_pii += pii
                if new_data != file_data:
                    modified_parts[idx] = new_data

            if file_count == 0:
                print(f"[PII] multipart: {flow.id} no file parts found", file=sys.stderr, flush=True)
                return data, 0

            # Rebuild multipart if any files were redacted
            if modified_parts and _REDACT_FILES:
                try:
                    data = _rebuild_multipart(data, boundary, modified_parts)
                    print(f"[PII] multipart: {flow.id} rebuilt with {len(modified_parts)} redacted file(s)",
                          flush=True)
                except Exception as e:
                    print(f"[PII] multipart rebuild error: {e}", file=sys.stderr, flush=True)

            return data, total_pii

        except Exception as e:
            print(f"[PII] multipart parse error: {e}", file=sys.stderr, flush=True)
            return data, 0

    # --- Single file upload (ChatGPT oaiusercontent, etc.) ---
    total_pii, data = _process_one_file(flow.id, data, content_type, url)
    return data, total_pii


def _process_one_file(flow_id: str, data: bytes, content_type: str, url: str,
                      filename: str = "", index: int = 1) -> tuple[int, bytes]:
    """Extract, scan, optionally redact a single file. Returns (pii_count, modified_bytes)."""
    size = len(data)
    label = f"{flow_id}[{index}]" if index > 1 else flow_id
    tag = f" (multipart[{index}] {filename})" if filename else ""

    # Skip CORS preflight / zero-byte requests — nothing to scan, avoid [unsupported: ] noise
    if size == 0:
        print(f"[PII] file: {label}{tag} empty body — skipped (preflight/zero-byte)",
              file=sys.stderr, flush=True)
        return 0, data

    # Resolve ChatGPT filename from the create-file API registry (blob PUT has no filename)
    if not filename:
        m = re.search(r"oaiusercontent\.com/files/([^/]+)/raw", url)
        if m:
            meta = _upload_meta.get(m.group(1))
            if meta and meta.get("filename"):
                filename = meta["filename"]
                tag = f" (multipart[{index}] {filename})"

    # Sniff content type from magic bytes when the header is missing/empty
    if not content_type or content_type.lower() in ("application/octet-stream",):
        sniffed = _sniff_content_type(data)
        if sniffed != content_type:
            print(f"[PII] file: {label}{tag} header ct='{content_type}' -> sniffed '{sniffed}'",
                  file=sys.stderr, flush=True)
            content_type = sniffed

    info = {
        "flow_id": flow_id,
        "file_index": index,
        "filename": filename,
        "url": url[:200],
        "content_type": content_type,
        "size": size,
    }

    # Extract text (timed)
    t0 = datetime.now(timezone.utc)
    extracted = _extract_text(data, content_type)
    elapsed_ms = round((datetime.now(timezone.utc) - t0).total_seconds() * 1000)
    info["text_len"] = len(extracted)
    info["extract_ms"] = elapsed_ms
    print(f"[PII] file-extract: {label}{tag} {content_type} {size}B -> {len(extracted)} chars ({elapsed_ms}ms)",
          file=sys.stderr, flush=True)

    if not extracted.strip() or extracted.startswith("["):
        if extracted.startswith("["):
            info["error"] = extracted
        _log_file_event(info)
        print(f"[PII] file: {label}{tag} -> {extracted[:80]}",
              file=sys.stderr, flush=True)
        return 0, data

    # Preview for debugging
    preview = extracted[:200].replace("\n", "\\n")
    print(f"[PII] file-text: {label}{tag} -> \"{preview}\"", file=sys.stderr, flush=True)

    # Scan for PII
    _, store = _redact(extracted)

    if store:
        entities = list(store.values())
        info["pii_count"] = len(entities)
        info["pii_preview"] = [v[:80] for v in entities[:5]]
        print(f"[PII] file: {label}{tag} FOUND {len(entities)} PII -> {info['pii_preview']}",
              flush=True)

        if _REDACT_FILES:
            try:
                ct = content_type.lower()
                if "application/pdf" in ct:
                    data = _redact_pdf(data, store)
                elif "wordprocessingml" in ct or "msword" in ct:
                    data = _redact_docx(data, store)
                elif any(t in ct for t in ("image/jpeg", "image/png", "image/jpg", "image/webp")):
                    data = _redact_image(data, store)
                elif "text/plain" in ct:
                    for placeholder, original in store.items():
                        data = data.replace(original.encode(), placeholder.encode())
                info["redacted"] = True
                info["redacted_size"] = len(data)
            except Exception as e:
                info["redaction_error"] = str(e)
                print(f"[PII] file redaction error: {e}", file=sys.stderr, flush=True)
    else:
        print(f"[PII] file: {label}{tag} {size}B -> no PII ({info['text_len']} chars)",
              file=sys.stderr, flush=True)

    _log_file_event(info)
    return len(store), data


def _rebuild_multipart(original: bytes, boundary: str,
                       replacements: dict[int, bytes]) -> bytes:
    """Rebuild a multipart body with modified file payloads using the email library."""
    from email.mime.multipart import MIMEMultipart
    from email.mime.base import MIMEBase
    
    try:
        msg = BytesParser().parsebytes(
            ("Content-Type: multipart/form-data; boundary=" + boundary + "\r\n\r\n").encode()
            + original
        )
        parts = list(msg.walk())
        
        for idx in replacements:
            if idx < len(parts):
                parts[idx].set_payload(replacements[idx])
        
        # Serialize back
        out = io.BytesIO()
        # Generator flattens the multipart tree back to bytes
        gen = msg.as_bytes().split(b"\r\n\r\n", 1)
        if len(gen) > 1:
            return gen[1]
    except Exception:
        pass
    
    return original
def _dump_structure(obj, indent="", max_depth=4, _depth=0):
    """Recursively print JSON structure with string previews."""
    if _depth > max_depth:
        print(f"{indent}...", file=sys.stderr, flush=True)
        return
    if isinstance(obj, dict):
        for k, v in obj.items():
            if isinstance(v, str):
                preview = v[:200].replace("\n", "\\n")
                print(f"{indent}{k}: \"{preview}\"", file=sys.stderr, flush=True)
            elif isinstance(v, (list, dict)):
                print(f"{indent}{k}: [{_count(v)}]", file=sys.stderr, flush=True)
                _dump_structure(v, indent + "  ", max_depth, _depth + 1)
            else:
                print(f"{indent}{k}: {v}", file=sys.stderr, flush=True)
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            if isinstance(item, str):
                preview = item[:200].replace("\n", "\\n")
                print(f"{indent}[{i}]: \"{preview}\"", file=sys.stderr, flush=True)
            elif isinstance(item, (dict, list)):
                print(f"{indent}[{i}]:", file=sys.stderr, flush=True)
                _dump_structure(item, indent + "  ", max_depth, _depth + 1)
            else:
                print(f"{indent}[{i}]: {item}", file=sys.stderr, flush=True)


def _count(obj):
    """Return len(obj) as string, or '?'."""
    try:
        return str(len(obj))
    except Exception:
        return "?"


# ---------------------------------------------------------------------------
# Endpoint matcher --- only intercept known LLM API endpoints
# ---------------------------------------------------------------------------
ENDPOINTS = [
    (r"chatgpt\.com/backend-anon/f/conversation",       1),
    ("chatgpt.com/backend-api/f/conversation",           1),
    (r"chatgpt\.com/unauth-mweb/conversation/updates",   1),
    (r"claude\.ai/api/organizations/[^/]+/chat_conversations/[^/]+/completion", 1),
    ("api.anthropic.com/v1/messages",                    1),
    ("/v1/chat/completions",                             1),
]

# ---------------------------------------------------------------------------
# File upload endpoints — intercept raw file bytes before they reach the LLM
# ---------------------------------------------------------------------------
FILE_UPLOAD_ENDPOINTS = [
    # ChatGPT: PUT to Azure blob storage (SAS-authenticated)
    (r"oaiusercontent\.com/files/.*/raw", 1),
    # Claude: multipart upload
    (r"claude\.ai/api/organizations/[^/]+/.*upload-file", 1),
    (r"claude\.ai/api/organizations/[^/]+/convert_document", 1),
]

# Gate: set PII_REDACT_FILES=1 to modify file bodies (best-effort)
_REDACT_FILES = os.environ.get("PII_REDACT_FILES", "0") == "1"

# CRLF constant — avoids \\r\\n escape mangling across tools
_CRLF = "\r\n"

# Debug: dump raw request body (set PII_DEBUG=1 to enable)
_PII_DEBUG = os.environ.get("PII_DEBUG", "0") == "1"

# Thread-local flow ID
_flow = threading.local()

# ChatGPT filename registry: maps oaiusercontent blob uuid -> {"filename", "content_type"}.
# Populated by intercepting the create-file API response; consumed by the blob PUT.
_upload_meta: dict[str, dict] = {}
_pending_filenames: dict[str, str] = {}  # flow.id -> filename (from create-file multipart)


def flow_id() -> str:
    return getattr(_flow, "id", "")


# ---------------------------------------------------------------------------
# mitmproxy handlers
# ---------------------------------------------------------------------------
def _extract_multipart_filename(flow: http.HTTPFlow) -> str:
    """Extract the first filename= from a multipart request body (create-file API)."""
    try:
        ct = flow.request.headers.get("content-type", "")
        boundary = None
        for part in ct.split(";"):
            p = part.strip()
            if p.startswith("boundary="):
                boundary = p.split("=", 1)[1].strip("\"'")
                break
        if not boundary:
            return ""
        header = "Content-Type: " + ct + _CRLF + _CRLF
        msg = BytesParser().parsebytes(header.encode() + (flow.request.content or b""))
        for part_msg in msg.walk():
            disp = str(part_msg.get("Content-Disposition", ""))
            if "filename=" in disp:
                for token in disp.split(";"):
                    token = token.strip()
                    if token.startswith("filename="):
                        return token.split("=", 1)[1].strip("\"'")
        return ""
    except Exception:
        return ""


def request(flow: http.HTTPFlow):
    """Intercept outgoing requests to LLM APIs, redact PII from bodies."""
    url = flow.request.pretty_url

    # --- ChatGPT create-file API: capture filename, correlate blob PUT later ---
    if re.search(r"chatgpt\.com/(backend-api|backend-anon)/files", url):
        _flow.id = flow.id
        fname = _extract_multipart_filename(flow)
        if fname:
            _pending_filenames[flow.id] = fname
            print(f"[PII] files-api {flow.id} filename={fname}", file=sys.stderr, flush=True)
        return

    # --- File upload check (runs before text endpoint matching) ---
    file_matched = any(re.search(p, url) for p, _ in FILE_UPLOAD_ENDPOINTS)
    if file_matched:
        _flow.id = flow.id
        ct = flow.request.headers.get("content-type", "")
        print(f"[PII] file-upload {flow.id} ct={ct[:40]}", file=sys.stderr, flush=True)
        try:
            new_body, count = _scan_file_upload(flow)
            if count > 0 and new_body != flow.request.content:
                flow.request.content = new_body
                print(f"[PII] file redacted {count} PII in {flow.id}", flush=True)
        except Exception as e:
            print(f"[PII] file ERROR in {flow.id}: {e}", file=sys.stderr, flush=True)
        return

    # --- Text endpoint check ---
    matched = any(re.search(p, url) for p, _ in ENDPOINTS)
    if not matched:
        return

    _flow.id = flow.id
    content_type = flow.request.headers.get("content-type", "")

    # Always log URL + content-type for debugging
    path = url.split("chatgpt.com", 1)[-1] if "chatgpt.com" in url else url
    print(f"[PII] {flow.id} ct={content_type[:40]} url={path[:120]}", file=sys.stderr, flush=True)

    try:
        # --- Form-encoded (ChatGPT unauth-mweb) ---
        if "application/x-www-form-urlencoded" in content_type:
            new_body, count = _scan_form(flow.request.content, url, content_type)
            if count > 0:
                flow.request.content = new_body
                print(f"[PII] redacted {count} field(s) in {flow.id}", flush=True)
            else:
                print(f"[PII] form: no PII in {flow.id}", file=sys.stderr, flush=True)
            return

        # --- JSON body (ChatGPT backend-api, Claude, etc.) ---
        body = json.loads(flow.request.content)

        if _PII_DEBUG:
            print(f"[PII] DEBUG {flow.id} url={url}", file=sys.stderr, flush=True)
            _dump_structure(body, indent="  ")

        new_body, count = _scan_json(body, url, content_type)
        if count > 0:
            flow.request.content = new_body
            print(f"[PII] redacted {count} field(s) in {flow.id}", flush=True)
        else:
            print(f"[PII] scanned: no PII in {flow.id}", file=sys.stderr, flush=True)
    except json.JSONDecodeError:
        print(f"[PII] skip: non-JSON body in {flow.id}", file=sys.stderr, flush=True)
    except Exception as e:
        print(f"[PII] ERROR in request: {e}", file=sys.stderr, flush=True)


#def response(flow: http.HTTPFlow):
 #   """Restore original PII values in the API response."""
    # --- ChatGPT create-file API: map upload_url blob uuid -> filename ---
  #  if flow.id in _pending_filenames:
   #     filename = _pending_filenames.pop(flow.id)
    #    try:
     #       text = flow.response.get_text() or "{}"
      #      resp = json.loads(text)
       #     # The response contains the upload URL (and often the filename too)
        #    upload_url = resp.get("upload_url") or resp.get("url") or ""
         #   m = re.search(r"oaiusercontent\.com/files/([^/]+)/raw", upload_url)
          #  blob_uuid = m.group(1) if m else ""
           # if blob_uuid:
            #    _upload_meta[blob_uuid] = {
             #       "filename": filename,
              #      "content_type": resp.get("content_type") or resp.get("mime_type") or "",
               # }
                #print(f"[PII] files-api mapped {blob_uuid} -> {filename}",
                 #     file=sys.stderr, flush=True)
       # except Exception as e:
        #    print(f"[PII] files-api response error: {e}", file=sys.stderr, flush=True)
        #return

    #try:
     #   store = _pii_store.pop(flow.id, None)
      #  if store:
       #     text = flow.response.get_text()
        #    if text:
         #       for placeholder, original in store.items():
          #          text = text.replace(placeholder, original)
           #     flow.response.set_text(text)
            #    _log({"event": "restore", "placeholders": len(store)})
    #except Exception as e:
     #   print(f"[PII] ERROR in response: {e}", file=sys.stderr, flush=True)
